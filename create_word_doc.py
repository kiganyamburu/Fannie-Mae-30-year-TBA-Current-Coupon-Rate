"""
Script to create Word document for Homework Assignment #2 - Problems 1, 2, and 3
"""

# Install python-docx if not available
import subprocess
import sys

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
except ImportError:
    print("Installing python-docx...")
    subprocess.check_call([sys.executable, '-m', 'pip', 'install', 'python-docx', '-q'])
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE

# Create document
doc = Document()

# Title
title = doc.add_heading('Homework Assignment #2', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Subtitle
subtitle = doc.add_paragraph('Problems 1, 2, and 3')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()  # Spacer

# Student info section
doc.add_paragraph('Student Name: _______________________________')
doc.add_paragraph('Date: _______________________________')
doc.add_paragraph()

# ============================================================================
# PROBLEM 1
# ============================================================================
doc.add_heading('Problem 1: Urban Institute Chart Book Review (0.25 points)', level=1)

doc.add_paragraph(
    'Review the December 2025 Chart Book published by the Urban Institute. '
    'The Chart Book can be found at: https://www.urban.org/research/publication/housing-finance-policy-center-chartbook'
)
doc.add_paragraph()

# Question 2-1
doc.add_heading('Question 2-1: Top Five New Things Learned', level=2)
doc.add_paragraph(
    'Highlight the top five things that you learned new from the chart book:'
)

for i in range(1, 6):
    doc.add_paragraph(f'{i}. ', style='List Number')
    doc.add_paragraph()  # Space for answer

doc.add_paragraph()

# Question 2-2
doc.add_heading('Question 2-2: Three Things Not Understood', level=2)
doc.add_paragraph(
    'State three things that you do not understand from the chart book:'
)

for i in range(1, 4):
    doc.add_paragraph(f'{i}. ', style='List Number')
    doc.add_paragraph()  # Space for answer

doc.add_page_break()

# ============================================================================
# PROBLEM 2
# ============================================================================
doc.add_heading('Problem 2: Recursion Glossary Review (0.25 points)', level=1)

doc.add_paragraph(
    'Review Fannie Mae loan level terms from the Recursion Glossary, and list the terms that you have confusion about.'
)
doc.add_paragraph()

doc.add_heading('Common Loan-Level Terms for Reference:', level=2)

terms_table = doc.add_table(rows=1, cols=2)
terms_table.style = 'Table Grid'
hdr_cells = terms_table.rows[0].cells
hdr_cells[0].text = 'Term'
hdr_cells[1].text = 'Description'

common_terms = [
    ('LTV', 'Loan-to-Value Ratio'),
    ('CLTV', 'Combined Loan-to-Value Ratio'),
    ('DTI', 'Debt-to-Income Ratio'),
    ('FICO', 'Credit Score'),
    ('UPB', 'Unpaid Principal Balance'),
    ('CPR', 'Conditional Prepayment Rate'),
    ('SMM', 'Single Monthly Mortality'),
    ('WAC', 'Weighted Average Coupon'),
    ('WAM', 'Weighted Average Maturity'),
    ('WALA', 'Weighted Average Loan Age'),
]

for term, desc in common_terms:
    row_cells = terms_table.add_row().cells
    row_cells[0].text = term
    row_cells[1].text = desc

doc.add_paragraph()

doc.add_heading('Terms That Cause Confusion:', level=2)
doc.add_paragraph(
    'List the Fannie Mae loan-level terms from the Recursion Glossary that you find confusing or need clarification:'
)

for i in range(1, 11):
    p = doc.add_paragraph(style='List Number')
    p.add_run(f'Term: ').bold = True
    p.add_run('_______________________________')
    doc.add_paragraph('   Reason for confusion: _______________________________')
    doc.add_paragraph()

doc.add_page_break()

# ============================================================================
# PROBLEM 3
# ============================================================================
doc.add_heading('Problem 3: Agency Primary Mortgage Market Size (0.5 points)', level=1)

doc.add_paragraph(
    'Using Recursion\'s Cohort Analyzer, analyze the Agency primary mortgage market size '
    'from January 2025 to December 2025.'
)
doc.add_paragraph()

# Definitions box
doc.add_heading('Key Definitions:', level=2)
defs = doc.add_paragraph()
defs.add_run('GSE: ').bold = True
defs.add_run('Government-Sponsored Enterprises (Fannie Mae and Freddie Mac)\n')
defs.add_run('GNM: ').bold = True
defs.add_run('Ginnie Mae (Government National Mortgage Association)\n')
defs.add_run('Agency: ').bold = True
defs.add_run('Fannie Mae, Freddie Mac, and Ginnie Mae combined')

doc.add_paragraph()

# Question 3-1
doc.add_heading('Question 3-1: Total Outstanding Balances by Agency (Monthly)', level=2)
doc.add_paragraph(
    'Report the total outstanding balances of loans for each month by each Agency:'
)

# Create table for outstanding balances
balance_table = doc.add_table(rows=13, cols=4)
balance_table.style = 'Table Grid'

# Header row
headers = ['Month', 'Fannie Mae ($B)', 'Freddie Mac ($B)', 'Ginnie Mae ($B)']
for i, header in enumerate(headers):
    balance_table.rows[0].cells[i].text = header

# Month rows
months = ['Jan 2025', 'Feb 2025', 'Mar 2025', 'Apr 2025', 'May 2025', 'Jun 2025',
          'Jul 2025', 'Aug 2025', 'Sep 2025', 'Oct 2025', 'Nov 2025', 'Dec 2025']

for i, month in enumerate(months, 1):
    balance_table.rows[i].cells[0].text = month

doc.add_paragraph()

# Question 3-2
doc.add_heading('Question 3-2: Total Issuance Volumes by Agency (Monthly)', level=2)
doc.add_paragraph(
    'Report the total issuance volumes of loans issued into Ginnie Mae, Fannie Mae, '
    'and Freddie Mac for each month of year 2025:'
)

# Create table for issuance volumes
issuance_table = doc.add_table(rows=13, cols=4)
issuance_table.style = 'Table Grid'

# Header row
for i, header in enumerate(headers):
    issuance_table.rows[0].cells[i].text = header

# Month rows
for i, month in enumerate(months, 1):
    issuance_table.rows[i].cells[0].text = month

doc.add_paragraph()

# Analysis section
doc.add_heading('Analysis and Observations:', level=2)
doc.add_paragraph(
    'Provide your analysis of the trends observed in outstanding balances and issuance volumes:'
)
doc.add_paragraph()
doc.add_paragraph('_' * 80)
doc.add_paragraph('_' * 80)
doc.add_paragraph('_' * 80)
doc.add_paragraph('_' * 80)
doc.add_paragraph('_' * 80)

# Save the document
output_path = 'Homework_Assignment_2_Problems_1_2_3.docx'
doc.save(output_path)
print(f"✓ Word document created successfully: {output_path}")
